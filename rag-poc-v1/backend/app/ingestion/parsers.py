import os
from pypdf import PdfReader
import docx
from app.core.logging import logger

def parse_pdf(file_path: str) -> str:
    try:
        reader = PdfReader(file_path)
        pages_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)
        
        full_text = "\n\n--- PAGE BREAK ---\n\n".join(pages_text)
        
        if len(full_text.strip()) < 200:
            logger.warning(f"Likely scanned PDF: {os.path.basename(file_path)}; OCR not implemented in v1.")
            
        return full_text
    except Exception as e:
        logger.error(f"Failed to parse PDF {file_path}: {e}")
        return ""

def parse_docx(file_path: str) -> str:
    try:
        doc = docx.Document(file_path)
        content_items = []

        # Attempt sequential block iteration if available (python-docx >= 1.1.0)
        iter_inner = getattr(doc, "iter_inner_content", None)
        
        if callable(iter_inner):
            from docx.text.paragraph import Paragraph
            from docx.table import Table
            
            for block in iter_inner():
                if isinstance(block, Paragraph):
                    text = block.text.strip()
                    if text:
                        content_items.append(text)
                elif isinstance(block, Table):
                    for row in block.rows:
                        row_data = []
                        for cell in row.cells:
                            cell_text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                            if cell_text:
                                row_data.append(cell_text)
                        if row_data:
                            content_items.append("TABLE: " + " | ".join(row_data))
        else:
            # Safe Fallback for older python-docx: Paragaphs first, then Tables
            for p in doc.paragraphs:
                text = p.text.strip()
                if text:
                    content_items.append(text)
                    
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                        if cell_text:
                            row_data.append(cell_text)
                    if row_data:
                        content_items.append("TABLE: " + " | ".join(row_data))

        return "\n\n".join(content_items)
    except Exception as e:
        logger.error(f"Failed to parse DOCX {file_path}: {e}")
        return ""

def parse_file(file_path: str, doc_type: str) -> str:
    if doc_type == 'pdf':
        return parse_pdf(file_path)
    elif doc_type == 'docx':
        return parse_docx(file_path)
    return ""
